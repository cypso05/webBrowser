# Import required modules
import sys
import os
import time
import requests
from PyQt5.QtCore import Qt, QUrl, QEvent, QTimer
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QToolBar, QLineEdit, QAction,
    QVBoxLayout, QWidget, QFileDialog, QInputDialog, QMenu, QTabWidget, QPushButton, QMessageBox
)
from PyQt5.QtWebEngineWidgets import QWebEngineView, QWebEnginePage
from PyQt5.QtGui import QImage, QPainter
from PyQt5.QtPrintSupport import QPrinter
from docx import Document
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, QMessageBox
from PyQt5.QtGui import QTextDocument
import re

# Browser main window class
class Browser(QMainWindow):
    # Initialize UI elements
    def __init__(self):
        super().__init__()
        # Set window title
        self.setWindowTitle("Web Browser")
         # Create tabs widget
        self.tabs = QTabWidget()
        self.tabs.setTabsClosable(True)
        # Connect tab close signal to close tab method
        self.tabs.tabCloseRequested.connect(self.close_tab)
        
        # Create URL bar

        self.url_bar = QLineEdit()
        self.url_bar.returnPressed.connect(self.navigate_to_url)
        
        # Create toolbar buttons

        self.back_button = self.create_action("Back", self.navigate_back)
        self.forward_button = self.create_action("Forward", self.navigate_forward)
        self.reload_button = self.create_action("Reload", self.reload_page)
        self.home_button = self.create_action("Home", self.go_home)
        self.extract_content_button = self.create_action("Extract Content", self.extract_content)
        self.save_pdf_button = self.create_action("Save as PDF", self.save_as_pdf)
        self.save_word_button = self.create_action("Save as Word", self.save_as_word)
        self.search_keyword_button = self.create_action("Search Keywords", self.search_keywords)
        self.download_button = self.create_action("Download Content", self.download_file)
        
        # Add buttons to toolbar

        self.toolbar = QToolBar()
        self.add_actions_to_toolbar()
        self.create_new_tab_button()
        
        # Set main layout

        self.layout = QVBoxLayout()
        self.layout.addWidget(self.toolbar)
        self.layout.addWidget(self.tabs)
        # Create central widget
        container = QWidget()
        container.setLayout(self.layout)
        self.setCentralWidget(container)
        # Connect signal to update URL on tab change
        self.tabs.currentChanged.connect(self.on_current_tab_changed)

        self.link_hovered = False
        # Open home page tab on startup
        self.create_new_tab()
         # Remove margins around layout
        self.layout.setContentsMargins(0, 0, 0, 0)
    # Helper method to create toolbar buttons
    def create_action(self, text, slot):
        action = QAction(text, self)
        action.triggered.connect(slot)
        return action
    # Add buttons to toolbar
    def add_actions_to_toolbar(self):
        actions = [
            self.back_button, self.forward_button, self.reload_button,
            self.home_button, self.extract_content_button, self.save_pdf_button,
            self.save_word_button, self.search_keyword_button, self.download_button
        ]
        for action in actions:
            self.toolbar.addAction(action)
             # Add separator
        self.toolbar.addSeparator()
         # Add URL bar
        self.toolbar.addWidget(self.url_bar)
    
    # Update URL bar on tab change
    def on_current_tab_changed(self, index):
        qurl = self.tabs.currentWidget().url() if self.tabs.currentWidget() else QUrl()
        self.update_urlbar(qurl)
 # Set URL bar text
    def update_urlbar(self, qurl):
        self.url_bar.setText(qurl.toString())
        self.url_bar.setCursorPosition(0)

    def create_new_tab_button(self):
        new_tab_button = QPushButton("+", self)
         # Connect clicked signal
        new_tab_button.clicked.connect(self.create_new_tab)
        self.toolbar.addWidget(new_tab_button)

    def create_new_tab(self, url="https://www.google.com"):
        # Create browser view 
        browser = QWebEngineView()
        # Set custom web page
        browser.setPage(BrowserWebEnginePage(browser))

        browser.page().linkHovered.connect(self.on_link_hovered)
        browser.customContextMenuRequested.connect(self.open_link_context_menu)
        # Install event filter to handle clicks
        browser.installEventFilter(self)
        
         # Add JavaScript to handle fullscreen video

        javascript_code = """
            var videoPlayer = document.querySelector('video');
            if (videoPlayer) {
                videoPlayer.addEventListener('webkitfullscreenchange', function() {
                    if (document.webkitIsFullScreen) {
                        document.webkitExitFullscreen();
                    } else {
                        videoPlayer.webkitRequestFullscreen();
                    }
                });
            }
        """
        browser.page().runJavaScript(javascript_code)
        # Connect signals to update on URL change and page load
         
        browser.urlChanged.connect(self.update_urlbar)
        browser.loadFinished.connect(self.update_title)
        # Load page  on new tab                                                     
        browser.setUrl(QUrl.fromUserInput(url) if url else QUrl('about:blank'))

        i = self.tabs.addTab(browser, "New Tab")
        self.tabs.setCurrentIndex(i)

    def close_tab(self, index):
         # Don't close last tab
        if self.tabs.count() < 2:
            return
         # Remove tab
        self.tabs.removeTab(index)
     # Navigate back in history and other navigations 
    def navigate_back(self):
        self.active_tab().back()

    def navigate_forward(self):
        self.active_tab().forward()

    def reload_page(self):
        self.active_tab().reload()

    def active_tab(self):
        return self.tabs.currentWidget()

    def navigate_to_url(self):
        q = QUrl(self.url_bar.text())
        if q.scheme() == "":
            q.setScheme("http")
        self.create_new_tab(q.toString())
    # Update window title to page title
    def update_title(self):
        index = self.tabs.currentIndex()
        title = self.active_tab().page().title()
        # Set tab title
        self.tabs.setTabText(index, title)
        # Set window title
        self.setWindowTitle(f"{title} - Web Browser")

    def go_home(self):
        self.active_tab().setUrl(QUrl("https://www.google.com"))

    def extract_content(self):
        self.active_tab().page().toHtml(self.process_page_html_for_media)

    def process_page_html_for_media(self, html):
        try:
            pattern = re.compile(r'(https?://[^\s]+\.(mp3|mp4|avi|mov|flv|wmv|mkv))', re.IGNORECASE)
            urls = pattern.findall(html)
            media_urls = list(set([url[0] for url in urls]))

            if not media_urls:
                QMessageBox.information(self, "Extract Content", "No media content found on this page.")
                return

            for media_url in media_urls:
                self.download_media(media_url)

        except Exception as e:
            QMessageBox.warning(self, "Error", f"An error occurred while processing media content: {str(e)}")

    def save_as_pdf(self):
        pdf_folder = os.path.join(os.path.expanduser("~"), "Desktop")
        pdf_path = os.path.join(pdf_folder, f"webpage_{time.strftime('%Y%m%d%H%M%S')}.pdf")

        self.take_screenshot_and_save_as_pdf(pdf_path)

    def take_screenshot_and_save_as_pdf(self, pdf_path):
        active_tab = self.active_tab()

        # Take HTML content of the entire page
        active_tab.page().toHtml(self.take_screenshot_callback(pdf_path))

    def take_screenshot_callback(self, pdf_path):
        def callback(html):
            active_tab = self.active_tab()

            # Use QTextDocument to render HTML content
            document = QTextDocument()
            document.setHtml(html)

            # Save the document as a PDF
            pdf_printer = QPrinter(QPrinter.HighResolution)
            pdf_printer.setOutputFormat(QPrinter.PdfFormat)
            pdf_printer.setOutputFileName(pdf_path)

            document.print_(pdf_printer)

        return callback

    def save_as_word(self):
        word_path, _ = QFileDialog.getSaveFileName(self, "Save Word", "", "Word Files (*.docx)")
        if word_path:
            self.active_tab().page().toPlainText(lambda text: self.write_to_word(text, word_path))

    def write_to_word(self, text, word_path):
        document = Document()
        document.add_paragraph(text)
        document.save(word_path)

    def search_keywords(self):
        keyword, ok = QInputDialog.getText(self, "Search Keywords", "Enter keyword:")
        if ok and keyword:
            self.active_tab().page().findText(keyword, QWebEnginePage.FindFlags(0), self.keyword_search_callback)

    def keyword_search_callback(self, result):
        QMessageBox.information(self, "Keyword Search", f"Occurrences: {result}")

    def download_file(self):
        url, ok = QInputDialog.getText(self, "Download File", "Enter the URL of the file to download:")
        if ok and url:
            file_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "All Files (*)")
            if file_path:
                self.download_file_by_url(url, file_path)

    def download_file_by_url(self, url, file_path):
        response = requests.get(url)
        if response.status_code == 200:
            with open(file_path, 'wb') as file:
                file.write(response.content)
            QMessageBox.information(self, "Download Complete", "Download successful.")
        else:
            QMessageBox.warning(self, "Download Failed", "Failed to download the file.")
            
    # Handle link hover
    def on_link_hovered(self, url):
        self.link_hovered = True
        
        
# Handle right-click menu on links  

    def open_link_context_menu(self, pos):
        context_menu = QMenu(self)
        open_in_new_tab_action = context_menu.addAction("Open Link in New Tab")
        open_in_new_window_action = context_menu.addAction("Open Link in New Window")
        
          # Get action

        action = context_menu.exec_(self.mapToGlobal(pos))
        
         # Open in tab or window

        if action == open_in_new_tab_action:
            url = self.active_tab().page().url().resolvedUrl().toString()
            self.create_new_tab(url)
        elif action == open_in_new_window_action:
            url = self.active_tab().page().url().resolvedUrl().toString()
            self.create_new_window(url)

        self.link_hovered = False
# Open URL in new window 
    def create_new_window(self, url):
        new_window = Browser()
        new_window.create_new_tab(url)
        new_window.showMaximized()
        
        # Handle middle click to open in new tab  
    def eventFilter(self, obj, event):
        if obj is self.active_tab() and event.type() == QEvent.MouseButtonPress:
            if event.button() == Qt.MiddleButton:
                self.create_new_tab(self.active_tab().page().url().resolvedUrl().toString())
                return True
        return super().eventFilter(obj, event)
    # Browser page class

class BrowserWebEnginePage(QWebEnginePage):
    def createWindow(self, _type):
        if _type == QWebEnginePage.WebBrowserTab:
            browser.create_new_tab()
            return browser.tabs.currentWidget().page()
        return super().createWindow(_type)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setApplicationName("Web Browser")
    browser = Browser()
    browser.showMaximized()
    sys.exit(app.exec_())
